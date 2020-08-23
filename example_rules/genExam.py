mark_10 = ['AQ2.docx','AQ3.docx','AQ4.docx','BQ2.docx','BQ3.docx','BQ4.docx']
mark_20 = ['AQ1.docx','AQ5.docx','AQ6.docx','BQ1.docx','BQ5.docx','BQ6.docx']

#=== This section generates a combination of sets in a variable called listA
from itertools import combinations
import numpy as np

paper_10=list(combinations(mark_10, 2))
paper_20=list(combinations(mark_20, 2))


i=0
A=0
paperA=[]
while i<15:
    j=0
    while j<15:
        paperA.append(list(paper_10[i])+(list(paper_20[j])))
        #print(paperA)
        A+=1
        j+=1
    i+=1
#=== End pf section 
#=== This section Removes identifies and makes an array of bad sets
badSet=[]
numSet=0
for sets in paperA:
    numQn=0
    if sets.count('AQ1.docx')==1:
        if sets.count('BQ1.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ2.docx')==1:
        if sets.count('BQ2.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ3.docx')==1:
        if sets.count('BQ3.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ4.docx')==1:
        if sets.count('BQ4.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ5.docx')==1:
        if sets.count('BQ5.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ6.docx')==1:
        if sets.count('BQ6.docx')==1:
            badSet.append(numSet)
    for questions in sets:
        numQn+=1
    numSet+=1

with open('badSet', 'w') as bad:
    bad.write("Total number of bad sets is: "+str(len(badSet))+"\n")
    for element in badSet:
        bad.write(str(paperA[element])+"\n")
#=== End pf section 
#=== This section Removes the bad sets from paperA and makes paperA2

paperA2=paperA

numRemove=0
for number in badSet:
    print(number)
    print(paperA2[number-numRemove])
    paperA2.remove(paperA2[number-numRemove])
    numRemove+=1

badSet=[]
numSet=0
for sets in paperA2:
    numQn=0
    if sets.count('AQ1.docx')==1:
        if sets.count('BQ.docx1')==1:
            badSet.append(numSet)
    if sets.count('AQ2.docx')==1:
        if sets.count('BQ2.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ3.docx')==1:
        if sets.count('BQ3.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ4.docx')==1:
        if sets.count('BQ4.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ5.docx')==1:
        if sets.count('BQ5.docx')==1:
            badSet.append(numSet)
    if sets.count('AQ6.docx')==1:
        if sets.count('BQ6.docx')==1:
            badSet.append(numSet)


#=== End pf section 
#=== This section composes a document with on input of 4 files using a cover page and end of paper page 

def press(file1, file2, file3, file4):
    from docxcompose.composer import Composer
    import docx
    from docx import Document as Document_compose
    file_cover=docx.Document('cover.docx')
    file_cover.add_page_break()
    file_cover.save('dummy.docx')
    master = Document_compose('dummy.docx')
    composer = Composer(master)
    
    
    filec1=docx.Document(file1)
    filec1.add_page_break()
    filec1.save('dummy1.docx')
    doc1 = Document_compose('dummy1.docx')
    composer.append(doc1)
    
    filec2=docx.Document(file2)
    filec2.add_page_break()
    filec2.save('dummy2.docx')
    doc2 = Document_compose('dummy2.docx')
    composer.append(doc2)
    
    filec3=docx.Document(file3)
    filec3.add_page_break()
    filec3.save('dummy3.docx')
    doc3 = Document_compose('dummy3.docx')
    composer.append(doc3)
    
    filec4=docx.Document(file4)
    filec4.add_page_break()
    filec4.save('dummy4.docx')
    doc4 = Document_compose('dummy4.docx')
    composer.append(doc4)
    
    doc5 = Document_compose('end.docx')
    composer.append(doc5)
    
    return composer




#=== End pf section 
#=== This section takes the input from paperA2 array and results in a paper {int}.docx 

num=0
with open('logfile', 'w') as log:
    log.write("Paper name\tQuestions\n")
    for things in list(paperA2):
        log.write(str(num)+'\t'+str(things)+"\n")
        press(str(things[0]),str(things[1]), str(things[2]), str(things[3])).save(str(num)+'.docx')
        num=num+1

#=== Removing the dummy files
import os
for num in ["1","2","3","4"]:
    os.remove("dummy"+num+".docx")
os.remove("dummy.docx")

